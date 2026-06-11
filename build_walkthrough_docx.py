"""Generate the self-guided (self-serve) walkthrough as a Word document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

NAVY = RGBColor(0x1E, 0x3A, 0x5F)
BLUE = RGBColor(0x25, 0x63, 0xEB)
GRAY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x15, 0x80, 0x3D)
PURPLE = RGBColor(0x6D, 0x28, 0xD9)

URL = "https://web-production-efe30.up.railway.app"
PASSWORD = "coastal2026"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def heading(text, size=15, color=NAVY, space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def labeled(label, text, lab_color, italic=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    l = p.add_run(label + "  "); l.bold = True; l.font.color.rgb = lab_color
    r = p.add_run(text); r.italic = italic
    return p

def nav(text):  return labeled("➤ Click:", text, NAVY)
def see(text):  return labeled("See:", text, GRAY)
def why(text):  return labeled("Why it matters:", text, GREEN)

def body(text, space_after=8):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    p.add_run(text); return p

# ── Title ──
t = doc.add_paragraph()
r = t.add_run("Live Intake and Referral Agent"); r.bold = True; r.font.size = Pt(23); r.font.color.rgb = NAVY
sub = doc.add_paragraph()
rs = sub.add_run("A self-guided walkthrough — explore the working agent yourself in ~15 minutes."); rs.italic = True; rs.font.size = Pt(12); rs.font.color.rgb = GRAY
doc.add_paragraph().add_run("—" * 30).font.color.rgb = RGBColor(0xCB,0xD5,0xE1)

# ── Access ──
heading("Getting in", 13)
ab = doc.add_paragraph(); ab.paragraph_format.space_after = Pt(4)
ab.add_run("Link:  ").bold = True
lk = ab.add_run(URL); lk.font.color.rgb = BLUE
ap = doc.add_paragraph(); ap.paragraph_format.space_after = Pt(8)
ap.add_run("Password:  ").bold = True
ap.add_run(PASSWORD)
body("Open the link, enter the password, and click “Enter.” You’ll land on the Hub. "
     "This is a live working agent on real referral packages — not slides. Take it at your own pace; "
     "nothing you click can break it.")

# ── click-path at a glance ──
heading("Where to click — at a glance", 13)
steps = [
    ("1", "Hub → read the cards (orientation)"),
    ("2", "Click the big “Live Intake and Referral Agent” card → “Start Live Processing”"),
    ("3", "Click the highlighted “James Holloway” row → step through with the “→” buttons"),
    ("4", "Back to Hub → “Your SOPs → The Agent’s Rules” → click “Ingest SOP → generate rules”"),
    ("5", "Hub → “Live Observability Dashboard”"),
    ("6", "Hub → “Referral Analytics”, then “Scale & Cost Calculator”"),
    ("7", "Hub → “8-Agent Value Chain” (the bigger picture)"),
]
gt = doc.add_table(rows=1, cols=2); gt.style = "Light Grid Accent 1"
gt.rows[0].cells[0].paragraphs[0].add_run("Stop").bold = True
gt.rows[0].cells[1].paragraphs[0].add_run("Where to click").bold = True
for s, w in steps:
    c = gt.add_row().cells; c[0].text = s; c[1].text = w

# ── 1 ──
heading("1.  The Hub — orientation")
nav("Nothing yet — just read the cards on the landing page.")
see("A grid of capability cards (Knowledge Graph, Your SOPs, Live Observability, Analytics, Scale & Cost, the 8-Agent Value Chain) and one large card to launch the live agent.")
why("This is one intake agent plus the layers around it — the rule engine it runs on, where its rules come from, how it’s observed, and how it scales. You’ll visit the key ones below.")

# ── 2 ──
heading("2.  Watch the agent work")
nav("The large “Live Intake and Referral Agent” card → then the green “Start Live Processing” button.")
see("A queue of real referrals processing in real time — patient, equipment, priority (with the reason), gaps, and outcome resolving row by row. Every row’s 3 source documents are openable (📄 📋 💊).")
why("It reads every page, extracts the fields, checks them against cited rules, scores its own confidence, and decides route-vs-flag — with no manual data entry. Outreach is drafted, never blind-sent.")
body("As rows process, the live activity log at the bottom of the screen cites the SOP clause behind "
     "each decision — e.g. “held per SOP §2.1”, “ICD conflict → clinical code per SOP §3.4”, “escalated "
     "per SOP §6.1 / §8.0”. That’s the agent applying your SOPs, live.")
body("Tip: let a few rows finish; the “James Holloway” row is highlighted for the next step.", space_after=10)

# ── 3 ──
heading("3.  Inside one referral — the “is it real?” moment")
nav("Click the highlighted James Holloway row (or “Open Pipeline”). Step through 01 → 07 with the “→” buttons at the bottom of each step.")
see("The full pipeline, stepped 01–07:  01 Channel Intake  ·  02 Data Normalization (every field traced to a source document)  ·  "
    "03 Intelligence Layer (cited knowledge-graph validation + an ICD-code conflict where the agent escalates instead of guessing, plus confidence tier)  ·  "
    "04 Gap Triage (what’s missing — hard-block vs soft)  ·  05 Digital Outreach (email, drafted)  ·  "
    "06 Voice Outreach (a live call, then the transcript)  ·  07 Route Output (the route-ready record).")
why("This is the credibility core: it shows its work at every step, and when evidence is split it stops and routes to a human. The restraint is the feature.")
body("Where SOPs show up here: on Step 03 (Intelligence Layer) you’ll see “Rules sourced from your SOPs →” "
     "(opens Stop 4); at Step 07 (Route Output) a “SOPs Applied to This Episode” card lists the exact SOP "
     "clauses that governed this referral — decision → clause → regulation. On Step 06 (Voice Outreach) the "
     "agent places a real call to a demo line (you don’t need to answer); the transcript appears after.", space_after=10)

# ── 4 ──
heading("4.  Your SOPs → the Agent’s Rules  (the differentiator)")
nav("Back to the Hub → “Your SOPs → The Agent’s Rules” → click “Ingest SOP → generate rules.”")
see("A sample intake SOP on the left; on click, the agent reads it and generates the structured rules live — each mapped to its SOP clause and the CMS authority behind it. Plus governance (versioned, owner, change-control).")
why("The agent doesn’t run on generic logic — it runs on YOUR procedures. Change an SOP, the rule regenerates under review. And in any referral’s audit trail, each decision traces back: decision → rule → SOP clause → regulation.")
body("Seeing it in the run: the live log cites the SOP behind each decision — e.g. “held per SOP §2.1”, "
    "“ICD conflict → clinical code per SOP §3.4”, “jurisdiction TX SLA 2d per SOP §5.2”, “escalated per "
    "SOP §6.1 / §8.0”. The agent loads your SOP rule set on every run and evaluates each rule against "
    "every referral — it’s not a fixed script.")
body("Where it is on the maturity curve (worth being precise): the agent reads and runs your SOP rules "
    "today and traces every decision to its clause and regulation. Having your SOP library also set the "
    "exact thresholds directly (confidence cut-offs, state SLAs) is the deployment-phase step — "
    "configured-from-your-SOPs today, deepening to generated-from-your-SOPs in deployment.", space_after=10)

# ── 5 ──
heading("5.  Live Observability — proof the guardrails work")
nav("Hub → “Live Observability Dashboard.”")
see("Live telemetry (throughput, latency, AI quality, knowledge-graph outcomes, cost) PLUS two things to look at: a per-referral detail table — every episode’s latency and factors individually, with outliers flagged — and a “Guardrail Verification” list of the cases the agent handled autonomously, no human, available for spot-check.")
why("It demonstrates the guardrails are working AND that they’re auditable: you can see exactly which cases the agent cleared on its own and open any one to verify the decision. Not just an average — every referral, inspectable.")
body("The live event log on this dashboard also cites each decision to its SOP clause (e.g. “per SOP §2.1”) — so you can watch your SOPs being applied as the agent runs.")

# ── 6 ──
heading("6.  Scale, cost & management view")
nav("Hub → “Referral Analytics”, then “Scale & Cost Calculator.”")
see("Analytics: auto-route and escalation rates, confidence tiers, jurisdiction mix, and the most common gaps. Scale & Cost: change the inputs and watch time and cost move — ~3¢ per referral (measured), throughput scaling with concurrency.")
why("Answers the two executive questions — “does it hold up at volume, and what does it cost?” — with numbers you can move yourself, and a management view of where work concentrates.")

# ── 7 ──
heading("7.  The bigger picture")
nav("Hub → “8-Agent Value Chain.”")
see("Referral-to-payment as eight modular agents; intake (what you just explored) is Agent 1, built and live. The rest are the roadmap, each deployable on its own.")
why("Frames today’s agent as the first step of a chain — not a one-off — deployable in your environment.")

# ── close ──
heading("In one line")
body("A working agent on real referral packages: it reads, validates against cited rules, knows its own "
     "confidence, escalates instead of guessing, runs off your SOPs, reaches out by email and phone, and "
     "logs every decision for audit — with the guardrails visible and inspectable. The natural next step is a "
     "48-hour proof on a sample of your own intake packages — no integration, nothing sensitive shared.")

heading("A few tips while you explore", 13)
for tip in [
    "Click “Start Live Processing” before expecting the queue or observability to fill — the screen updates during a run.",
    "Any completed row opens its full pipeline (same view as the Holloway example).",
    "The Voice step dials a demo line; you don’t need to do anything — just watch the call status and transcript.",
    "Don’t refresh the live agent screen mid-run — it resets the queue.",
]:
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3); p.add_run(tip).font.size = Pt(10.5)

desktop = os.path.join(os.path.expanduser("~"), "Desktop")
out1 = os.path.join(desktop, "Live Intake and Referral Agent - Self-Guided Walkthrough.docx")
out2 = os.path.join(os.path.dirname(__file__), "deliverables", "self_guided_walkthrough.docx")
doc.save(out1); doc.save(out2)
print("Saved:", out1)
print("Saved:", out2)
