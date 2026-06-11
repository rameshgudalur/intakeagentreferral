"""Render the sample Intake SOP (markdown) as a Word document — so it looks like a real client SOP."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
import os, re

NAVY = RGBColor(0x1E, 0x3A, 0x5F); GRAY = RGBColor(0x55, 0x55, 0x55)
SRC = Path(__file__).parent / "sops" / "coastal_intake_sop.md"
lines = SRC.read_text(encoding="utf-8").splitlines()

doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

def add_runs(p, text):
    # split on **bold** and render bold segments
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if seg == "":
            continue
        r = p.add_run(seg)
        if i % 2 == 1:
            r.bold = True

for ln in lines:
    s = ln.rstrip()
    if not s:
        continue
    if s.startswith("# "):
        p = doc.add_paragraph(); r = p.add_run(s[2:]); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
    elif s.startswith("## "):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
        r = p.add_run(s[3:]); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY
    elif s.startswith("**") and s.endswith("**"):
        p = doc.add_paragraph(); r = p.add_run(s.strip("*")); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY
    else:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
        add_runs(p, s)

desktop = os.path.join(os.path.expanduser("~"), "Desktop")
out1 = os.path.join(desktop, "Coastal DME - Intake SOP (sample).docx")
out2 = os.path.join(os.path.dirname(__file__), "deliverables", "coastal_intake_sop.docx")
doc.save(out1); doc.save(out2)
print("Saved:", out1); print("Saved:", out2)
